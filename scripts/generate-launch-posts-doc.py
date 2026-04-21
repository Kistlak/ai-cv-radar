"""
Generates docs/launch-posts.docx.
Ready-to-paste launch posts for Hacker News (Show HN) and Indie Hackers,
plus a short how-to-post checklist for each platform.
One-shot script - safe to delete after use.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "docs" / "launch-posts.docx"

doc = Document()

s = doc.styles["Normal"]
s.font.name = "Calibri"
s.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)


def shade(paragraph, color_hex="F4F4F4"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _inline(p, text):
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                p.add_run(text[i:])
                break
            r = p.add_run(text[i + 2:end])
            r.bold = True
            i = end + 2
        elif text[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                p.add_run(text[i:])
                break
            r = p.add_run(text[i + 1:end])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            i = end + 1
        else:
            nxt = len(text)
            for marker in ("**", "`"):
                pos = text.find(marker, i)
                if pos != -1 and pos < nxt:
                    nxt = pos
            p.add_run(text[i:nxt])
            i = nxt


def para(text):
    p = doc.add_paragraph()
    _inline(p, text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    _inline(p, text)


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    _inline(p, text)


def post_block(text):
    """Render a copy-paste post body in a shaded block, preserving paragraph breaks."""
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        if line.strip():
            _inline(p, line)
        else:
            p.add_run(" ")
        shade(p, "F7F9FC")
    doc.add_paragraph()


def callout(text, label="TIP"):
    p = doc.add_paragraph()
    shade(p, "FFF8E1")
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x57, 0x00)
    _inline(p, text)


# ============================================================
# TITLE
# ============================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("AI CV Radar - Launch Posts")
r.bold = True
r.font.size = Pt(22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Ready-to-paste posts for Hacker News (Show HN) and Indie Hackers, with platform-specific notes")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# ============================================================
# BEFORE YOU POST
# ============================================================
h1("Before you post - 5 minute checklist")

numbered("Replace `[YOUR_LINK]` with your live URL (Vercel deploy or custom domain).")
numbered("Replace `[GITHUB_LINK]` with the public repo URL, or remove that line if the repo stays private.")
numbered("Make sure the landing page loads in under 2 seconds and the sign-up flow works on a fresh incognito window. First impressions on HN/IH are unforgiving.")
numbered("If `FALLBACK_*` keys are set, decide if you want to keep them on for launch traffic. A front-page Show HN can put 5,000+ visitors through your funnel in a day - your Anthropic bill will reflect that. Consider a hard daily spend limit in the Anthropic console first.")
numbered("Have a tab open on the post for the first 2-3 hours to reply to comments. Both communities reward founders who show up.")

callout("Best posting windows: HN is most active 8am-11am US Pacific on weekdays. IH is steadier but Tue-Thu mornings (US time) tend to get the most eyeballs.", label="TIMING")

doc.add_paragraph()

# ============================================================
# POST 1 - HACKER NEWS
# ============================================================
h1("Post 1. Hacker News (Show HN)")

para("**Where to post:** https://news.ycombinator.com/submit")
para("**Category:** Use the title prefix `Show HN:` - that routes it to the Show HN section automatically.")
para("**URL field:** Put your live app URL (`[YOUR_LINK]`). Put the body text below in the **text** field. Do not include both a URL and text body unless the body adds context the URL doesn't (it does here, so include both).")

doc.add_paragraph()
h2("Title (max 80 chars)")
post_block("Show HN: AI CV Radar - ranks jobs against your CV and auto-fills applications")

h2("Body")
post_block(
    "Hi HN, I built AI CV Radar after watching too many friends spend their evenings on the same job-application chore: read the JD, rewrite the CV, write yet another cover letter, fill the same Greenhouse form for the 40th time.\n"
    "\n"
    "What it does:\n"
    "\n"
    "- Upload your CV once. A PDF is text-extracted and parsed into structured data (skills, experience, education) by Claude.\n"
    "- Search across Remotive, Adzuna, and JSearch on the free tier. With an Apify token, LinkedIn / Indeed / Glassdoor unlock and Claude drives the search as an agentic loop - calling actors, reading results, refining the query when results are weak.\n"
    "- Each job gets three AI actions: a fit analysis (strengths, gaps, what to emphasize), a tailored cover letter, and an ATS-friendly Word .docx CV rewritten for that role.\n"
    "- A bundled Chromium extension fills Greenhouse, Lever, Ashby, and Workable forms (plus a heuristic fallback) using your stored profile. It never submits - you review every field and click submit yourself.\n"
    "\n"
    "Stack: Next.js 16, React 19, Supabase (Postgres + auth + storage), Drizzle, Anthropic Claude (Sonnet 4.6 for generation, Haiku 4.5 for scoring), `docx` for Word output, Manifest V3 for the extension.\n"
    "\n"
    "It's BYOK - you paste your own Anthropic key in Settings and it's encrypted at rest with AES before it touches the database. Optional Apify / Adzuna / RapidAPI keys unlock more sources.\n"
    "\n"
    "Live: [YOUR_LINK]\n"
    "Code: [GITHUB_LINK]\n"
    "\n"
    "Things I'd love feedback on:\n"
    "\n"
    "1. The Apify-driven agentic search vs straight deterministic adapters - is the extra latency worth the better recall on weird queries?\n"
    "2. Auto Apply UX - what's the right balance between \"fill everything I can\" and \"don't surprise the user\"?\n"
    "3. If you've shipped something that solved cover-letter fatigue differently, I'd love to hear how."
)

callout("HN values directness over polish. Don't oversell. The community is allergic to marketing copy and rewards genuine technical detail and honest open questions.", label="HN STYLE")

doc.add_paragraph()
h2("After you post")
bullet("Reply to **every** top-level comment in the first 4 hours, even one-liners. HN's ranking penalises inactive threads.")
bullet("If someone reports a bug, fix it live and reply with the commit hash. This single move has carried more Show HNs to the front page than anything else.")
bullet("Do **not** ask anyone to upvote - HN detects voting rings and will dead-list the post. Just share the link with friends and let them decide.")

doc.add_paragraph()

# ============================================================
# POST 2 - INDIE HACKERS
# ============================================================
h1("Post 2. Indie Hackers")

para("**Where to post:** https://www.indiehackers.com/post (choose the **Launched** group, or **Show IH** if your account is new).")
para("**Format:** IH allows full Markdown - headers, bold, lists, links all render. Lean into it. The body can be longer than HN; founder story is welcome.")

doc.add_paragraph()
h2("Title")
post_block("Launching AI CV Radar: turn your CV into a job-application autopilot (BYOK Claude + browser extension)")

h2("Body")
post_block(
    "Hey Indie Hackers, I'm shipping AI CV Radar this week after months of watching friends grind through the same job-search loop: read the JD, rewrite the CV, write yet another cover letter, fill the same Greenhouse form for the 40th time.\n"
    "\n"
    "**The wedge**\n"
    "\n"
    "Most \"AI job tools\" are job boards with AI bolted on. I built it the other way around: your CV is the source of truth, and every listing either fits it or doesn't. Claude scores each job 0-100 with a one-line reason, and on demand generates three things per role - a fit analysis (honest about gaps), a tailored cover letter, and an ATS-friendly Word CV rewritten for that specific role.\n"
    "\n"
    "**The bit I'm most excited about**\n"
    "\n"
    "A Manifest V3 Chromium extension that fills Greenhouse, Lever, Ashby, and Workable application forms using your stored profile. It never submits - you review every field and click submit yourself. (I lost two days to iframe quirks; turns out most career sites embed the ATS in an iframe and `chrome.scripting.executeScript` defaults to main-frame-only. Adding `allFrames: true` was the fix.)\n"
    "\n"
    "**Stack**\n"
    "\n"
    "Next.js 16 (App Router), React 19, Supabase (Postgres + auth + private storage), Drizzle ORM, Anthropic Claude (Sonnet 4.6 for generation, Haiku 4.5 for scoring), `docx` for Word output, Apify for the LinkedIn / Indeed / Glassdoor sources, Playwright for end-to-end smoke tests.\n"
    "\n"
    "**Pricing model**\n"
    "\n"
    "BYOK - users bring their own Anthropic key (encrypted at rest). I've also wired optional `FALLBACK_*` env vars so curious users without keys can still try the app on my dime, with the obvious cost-control caveats. Long-term plan is probably a small hosted tier with a margin on top of pass-through cost, but launching BYOK lets me ship without payment infra.\n"
    "\n"
    "**Live:** [YOUR_LINK]\n"
    "**Code:** [GITHUB_LINK]\n"
    "\n"
    "**Open questions for the IH crowd:**\n"
    "\n"
    "1. BYOK vs hosted-with-margin for v1 - which would you go with, and why?\n"
    "2. Anyone running a similar BYOK Claude product? How are you handling the cost-surprise vs UX-friction trade-off on the fallback keys?\n"
    "3. Worth doing a free tier of, say, 3 jobs scored before BYOK kicks in - or does that just attract tyre-kickers?\n"
    "\n"
    "Brutal feedback welcome."
)

callout("IH rewards founder vulnerability and revenue numbers. If you have any - even \"$0 MRR, 0 users\" with a credible plan to change that - say so. The community responds to honesty about where you actually are.", label="IH STYLE")

doc.add_paragraph()
h2("After you post")
bullet("Reply within an hour to anything that lands. IH's feed moves slower than HN, so a thoughtful 4-paragraph reply is normal and appreciated.")
bullet("Cross-link your HN post if it does well - IH founders love seeing other founders break through on HN.")
bullet("Follow up with a milestone post in 2-4 weeks: \"100 sign-ups, here's what broke.\" These often outperform the launch post itself.")

doc.add_paragraph()

# ============================================================
# OTHER PLACES TO CONSIDER
# ============================================================
h1("Bonus - other places worth a post")

bullet("**r/SideProject** (Reddit) - low-friction, friendly. Same body works; trim to 3 paragraphs.")
bullet("**r/cscareerquestions** or **r/jobs** - only if you frame it as \"I built this to scratch my own itch, would love feedback\" - never as a sales pitch. Reddit detects that immediately.")
bullet("**Product Hunt** - bigger production. Needs a hunter, prepared assets (gallery images, 1-line tagline, demo GIF), and ideally launches Tue-Thu. Worth doing **after** HN/IH so you have social proof to reference.")
bullet("**LinkedIn** - your own audience. A short post with a screenshot and a link converts well if you've ever posted about job-hunting before.")
bullet("**Twitter/X** - thread version. First tweet = the wedge. Last tweet = the link. Reply to your own thread with the demo GIF.")

doc.add_paragraph()

callout("Don't post everywhere on the same day. Stagger by 1-2 days so you can fix bugs reported on platform N before platform N+1 sees them.", label="ORDER")


OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Wrote {OUT}")
