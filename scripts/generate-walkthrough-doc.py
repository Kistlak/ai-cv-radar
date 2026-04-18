"""
Generates docs/ai-cv-radar-walkthrough.docx.
One-shot script — safe to delete after use.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "docs" / "ai-cv-radar-walkthrough.docx"

doc = Document()

# ------- base styles -------
for style_name in ("Normal",):
    s = doc.styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)


# ------- helpers -------
def shade_cell(paragraph, color_hex="F2F2F2"):
    """Add a light grey background to a paragraph (for code blocks)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def h3(text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def para(text):
    """Plain paragraph. Use **bold** and `code` inline markers."""
    p = doc.add_paragraph()
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                p.add_run(text[i:])
                break
            run = p.add_run(text[i + 2 : end])
            run.bold = True
            i = end + 2
        elif text[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                p.add_run(text[i:])
                break
            run = p.add_run(text[i + 1 : end])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            i = end + 1
        else:
            # find next marker
            nxt = len(text)
            for marker in ("**", "`"):
                pos = text.find(marker, i)
                if pos != -1 and pos < nxt:
                    nxt = pos
            p.add_run(text[i:nxt])
            i = nxt


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                p.add_run(text[i:])
                break
            run = p.add_run(text[i + 2 : end])
            run.bold = True
            i = end + 2
        elif text[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                p.add_run(text[i:])
                break
            run = p.add_run(text[i + 1 : end])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            i = end + 1
        else:
            nxt = len(text)
            for marker in ("**", "`"):
                pos = text.find(marker, i)
                if pos != -1 and pos < nxt:
                    nxt = pos
            p.add_run(text[i:nxt])
            i = nxt


def code(snippet, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for line in snippet.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        shade_cell(p, "F4F4F4")
    # spacer
    doc.add_paragraph()


# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AI CV Radar — How This App Works")
r.bold = True
r.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A plain-English walkthrough for someone new to Next.js")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# ============================================================
# SECTION 0 — Next.js cheat sheet
# ============================================================
h1("Before we start: a 60-second Next.js cheat-sheet")

para(
    "Next.js is a framework built on top of React. The whole app lives in the `app/` folder. "
    "A few conventions are enough to read almost any file in this project:"
)
bullet("Files named `page.tsx` become pages users visit. The URL mirrors the folder path.")
bullet("Files named `route.ts` are backend API endpoints — like a tiny Express route.")
bullet("Files named `layout.tsx` wrap pages with shared chrome (header, theme, etc.).")
bullet(
    "A file with `'use client'` at the top runs in the browser. Everything else runs on the server. "
    "You mark a file `'use client'` when it needs useState, event handlers, or browser APIs."
)
bullet(
    "Folders in parentheses like `(auth)` are route groups — they organize files but do not appear in the URL. "
    "So `app/(auth)/login/page.tsx` is visited at `/login`, not `/(auth)/login`."
)
bullet("Square brackets like `[id]` in a folder name mean a dynamic URL segment.")
bullet(
    "`middleware` runs on every request, before the page or route handler. "
    "In this project it lives in `proxy.ts` and is used to gate protected pages."
)

# ============================================================
# SECTION 1 — LOGIN
# ============================================================
h1("1. Login (passwordless magic link)")

para(
    "We use Supabase Auth in its simplest mode — magic links. "
    "The user types their email, we email them a link, they click it, and they are signed in. "
    "No password, no OAuth, no forms to manage."
)

h2("The login page")
para(
    "Location: `app/(auth)/login/page.tsx`. "
    "This is a client component (`'use client'` at the top) because it has a form with state — "
    "the email input, a loading flag, and a 'sent' flag. Anything that needs state or event handlers must run in the browser."
)

code(
    "'use client'\n"
    "\n"
    "export default function LoginPage() {\n"
    "  const [email, setEmail] = useState('')\n"
    "  const [sent, setSent] = useState(false)\n"
    "\n"
    "  async function handleSubmit(e) {\n"
    "    e.preventDefault()\n"
    "    const supabase = createClient()\n"
    "    await supabase.auth.signInWithOtp({\n"
    "      email,\n"
    "      options: { emailRedirectTo: `${APP_URL}/auth/callback` },\n"
    "    })\n"
    "    setSent(true)\n"
    "  }\n"
    "}",
    caption="app/(auth)/login/page.tsx (trimmed)",
)

para(
    "**Why this shape?** Supabase already ships a complete magic-link flow — we just call it. "
    "`signInWithOtp` triggers the email; the `emailRedirectTo` tells Supabase which URL the link should bring the user back to."
)

h2("The callback route")
para(
    "Location: `app/auth/callback/route.ts`. "
    "When the user clicks the link in their email, Supabase bounces them to `/auth/callback?code=...`. "
    "A `route.ts` file is a backend endpoint — this one exchanges the one-time code for a real session cookie."
)

code(
    "export async function GET(request) {\n"
    "  const code = new URL(request.url).searchParams.get('code')\n"
    "  if (code) {\n"
    "    const supabase = await createClient()\n"
    "    const { data } = await supabase.auth.exchangeCodeForSession(code)\n"
    "\n"
    "    // Ensure a profile row exists for new users (idempotent)\n"
    "    await db.insert(profiles)\n"
    "      .values({ id: data.user.id, email: data.user.email })\n"
    "      .onConflictDoNothing()\n"
    "\n"
    "    return NextResponse.redirect('/dashboard')\n"
    "  }\n"
    "}",
    caption="app/auth/callback/route.ts (trimmed)",
)

para(
    "**Why insert a profile here?** Supabase stores the auth user in its own private table. "
    "We want our own `profiles` table so we can reference users from `cvs`, `searches`, `user_api_keys`, etc. "
    "`onConflictDoNothing()` makes this safe to run on every login — new users get a row, returning users are ignored."
)

h2("Protecting pages with middleware")
para(
    "Location: `proxy.ts` (the file at the project root). "
    "Middleware runs on every request before any page or route handler. "
    "We use it to bounce unauthenticated users to `/login` if they try to visit a protected page."
)

code(
    "const isProtectedRoute =\n"
    "  request.nextUrl.pathname.startsWith('/dashboard') ||\n"
    "  request.nextUrl.pathname.startsWith('/cv') ||\n"
    "  request.nextUrl.pathname.startsWith('/search') ||\n"
    "  request.nextUrl.pathname.startsWith('/settings')\n"
    "\n"
    "if (isProtectedRoute && !user) {\n"
    "  return NextResponse.redirect('/login')\n"
    "}",
    caption="lib/supabase/middleware.ts (trimmed)",
)

para(
    "**Why middleware instead of checking in every page?** Writing `if (!user) redirect(...)` at the top of "
    "every protected page is repetitive and easy to forget. Middleware handles it once."
)

# ============================================================
# SECTION 2 — CV UPLOAD
# ============================================================
h1("2. CV upload")

para(
    "Goal: the user drops a PDF → we pull the text out, ask Claude to structure it "
    "(skills, experience, education), and save both the raw text and structured JSON to the database."
)

h2("The drag-and-drop form")
para(
    "Location: `components/cv-upload-form.tsx`. Client component because it needs drag events and state."
)

code(
    "export default function CvUploadForm() {\n"
    "  const [loading, setLoading] = useState(false)\n"
    "  const router = useRouter()\n"
    "\n"
    "  async function handleFile(file) {\n"
    "    if (file.type !== 'application/pdf') return toast.error('PDF only')\n"
    "    setLoading(true)\n"
    "    const formData = new FormData()\n"
    "    formData.append('file', file)\n"
    "    const res = await fetch('/api/cv/upload', { method: 'POST', body: formData })\n"
    "    if (res.ok) {\n"
    "      toast.success('CV uploaded and parsed!')\n"
    "      router.refresh()   // re-renders the server component so fresh data shows up\n"
    "    }\n"
    "    setLoading(false)\n"
    "  }\n"
    "}",
    caption="components/cv-upload-form.tsx (trimmed)",
)

para(
    "**Why `FormData` instead of JSON?** JSON cannot carry binary file bytes. "
    "`FormData` is the standard HTML file-upload format and `fetch` knows how to send it."
)
para(
    "**Why `router.refresh()`?** The dashboard shows your parsed CV. It is a server component — it fetched the data "
    "when the page loaded. `router.refresh()` tells Next.js to re-run that fetch so the new CV shows up immediately without a full page reload."
)

h2("The upload endpoint (the brain of the flow)")
para(
    "Location: `app/api/cv/upload/route.ts`. This is where the real work happens. It runs on the server."
)

bullet("1. Verify the user is logged in (Supabase `getUser()`)")
bullet("2. Pull the file out of the form data and verify it is a PDF")
bullet("3. Fetch + decrypt the user's Anthropic API key from the database")
bullet("4. Extract plain text from the PDF with `unpdf`")
bullet("5. Upload the original PDF to Supabase Storage at `{userId}/{timestamp}.pdf`")
bullet("6. Ask Claude to turn the raw text into structured JSON")
bullet("7. Mark any previous CV as inactive, insert the new one as active")

code(
    "// 5. Extract text\n"
    "const pdf = await getDocumentProxy(new Uint8Array(buffer))\n"
    "const { text } = await extractText(pdf, { mergePages: true })\n"
    "\n"
    "// 6. Upload PDF to Supabase Storage\n"
    "await supabase.storage.from('cvs').upload(filePath, buffer, ...)\n"
    "\n"
    "// 7. Ask Claude to structure the CV\n"
    "const message = await client.messages.create({\n"
    "  model: 'claude-sonnet-4-6',\n"
    "  messages: [{ role: 'user', content: 'Extract JSON { name, skills, experience, ... }: ' + rawText }],\n"
    "})\n"
    "\n"
    "// 8. Deactivate old CV, save new one\n"
    "await db.update(cvs).set({ isActive: false }).where(eq(cvs.userId, user.id))\n"
    "await db.insert(cvs).values({ userId, filePath, rawText, structured, isActive: true })",
    caption="app/api/cv/upload/route.ts (simplified)",
)

para(
    "**Why save both `rawText` and `structured`?** "
    "The structured JSON is perfect for showing the user ('here is your parsed CV'). "
    "But for job matching later, we feed the raw text into Claude so it can see nuance — phrasing, projects, tone — that a structured summary would lose."
)
para(
    "**Why only one active CV?** The search flow needs a single source of truth for 'what CV am I matching against right now'. "
    "Keeping old CVs but marking them inactive also means you do not lose history."
)

# ============================================================
# SECTION 3 — SEARCH + AGENTIC
# ============================================================
h1("3. Search (with the agentic workflow)")

h2("The problem we are solving")
para(
    "Job-board APIs like Apify are powerful, but they are 'dumb pipes': "
    "you hand them a query, they hand back results. If your query is off, the results are junk — "
    "and your code has no way to notice and try something else."
)
para(
    "When Claude Desktop searches jobs with the same Apify tools, the output is excellent. "
    "That is because Claude iterates: it tries a query, looks at the results, refines, tries another angle, and stops when the set is good. "
    "A human does the same thing intuitively. A fixed `fetch(apify, query)` call cannot."
)
para(
    "So for this app we put Claude in the driver's seat — it calls Apify itself, inspects results, refines, and returns a curated list. "
    "That is what 'agentic' means here."
)

h2("The flow end-to-end")

h3("Step 1 — The search form")
para(
    "Location: `components/search-form.tsx`. "
    "The user picks optional keywords, optional location (with typeahead), remote-only toggle, source toggles, and how many jobs they want. "
    "On submit it posts the form values to `/api/search`."
)

h3("Step 2 — The API route starts the job")
para(
    "Location: `app/api/search/route.ts`. "
    "This is deliberately quick. It creates a row in the `searches` table with status `running`, "
    "returns the search ID to the client immediately, and kicks the heavy work off in the background using Next.js's `after(...)` helper."
)

code(
    "const [search] = await db.insert(searches).values({\n"
    "  userId, cvId, query, location, remoteOnly, sources, maxResults,\n"
    "  status: 'running',\n"
    "}).returning()\n"
    "\n"
    "// Fire-and-forget: the browser does not wait for this\n"
    "after(async () => { await runSearch(search.id, user.id) })\n"
    "\n"
    "return NextResponse.json({ searchId: search.id }, { status: 202 })",
    caption="app/api/search/route.ts (trimmed)",
)

para(
    "**Why not wait for the search to finish in the HTTP response?** "
    "The full pipeline takes 60-90 seconds. No browser will wait that long. "
    "Instead we return a search ID instantly, the browser navigates to the results page, and that page polls a status endpoint until the row flips to `complete`."
)

h3("Step 3 — The background worker")
para(
    "Location: `lib/run-search.ts`. "
    "This is the orchestrator. It loads the search + the active CV + the user's API keys, then fans out to two paths in parallel:"
)
bullet("**Cheap path:** free-tier APIs like Remotive, Adzuna, JSearch. Fast, no cost, often low relevance.")
bullet("**Agentic path:** the Apify-backed sources (LinkedIn, Indeed, Glassdoor), driven by Claude.")

code(
    "const [cheapJobs, agenticJobs] = await Promise.all([\n"
    "  fetchAllSourcesMultiQuery(queries, ..., cheapSources),\n"
    "  runAgenticSearch({ cvText, userQuery, location, remoteOnly, maxResults, ... }),\n"
    "])\n"
    "const rawJobs = dedupeJobs([...cheapJobs, ...agenticJobs])\n"
    "const scoredJobs = await scoreJobs(rawJobs, cv.rawText, primaryQuery, keys.anthropicKey)",
    caption="lib/run-search.ts (trimmed)",
)

para(
    "**Why run both in parallel?** `Promise.all` starts both at the same time. "
    "The whole search takes only as long as the slower one (the agentic path), not the sum of both."
)
para(
    "**Why still score with Claude after the agent already curated?** "
    "The agent picks which jobs to include. The scoring pass ranks them against the CV. Two independent signals are more reliable than one."
)

h3("Step 4 — The agentic engine (the key piece)")
para(
    "Location: `lib/agentic-search.ts`. This is where Claude drives Apify itself."
)

para(
    "We call the Anthropic Messages API with two things attached:"
)
bullet("**The Apify MCP server** — a hosted service at `https://mcp.apify.com` that exposes Apify's actors (LinkedIn scraper, Indeed scraper, etc.) as tools Claude can call.")
bullet("**Our own `finalize_jobs` tool** — a local tool with a strict schema. Claude must call this exactly once at the end to return its curated list.")

para(
    "MCP (Model Context Protocol) is a standard for letting AI use external tools. "
    "The same plumbing that powers Claude Desktop's tool use powers this flow — which is why it works here as well as it does there."
)

code(
    "const response = await client.beta.messages.create({\n"
    "  model: 'claude-sonnet-4-6',\n"
    "  system: buildSystemPrompt(targetCount),\n"
    "  messages: [{ role: 'user', content: buildUserPrompt(input) }],\n"
    "\n"
    "  // Attach Apify's hosted MCP server — Anthropic handles the round-trips server-side\n"
    "  mcp_servers: [{\n"
    "    type: 'url',\n"
    "    url: 'https://mcp.apify.com',\n"
    "    name: 'apify',\n"
    "    authorization_token: input.apifyToken,\n"
    "  }],\n"
    "\n"
    "  tools: [\n"
    "    { type: 'mcp_toolset', mcp_server_name: 'apify' },   // Apify tools\n"
    "    FINALIZE_JOBS_TOOL,                                   // our local tool\n"
    "  ],\n"
    "  betas: ['mcp-client-2025-11-20'],\n"
    "})",
    caption="lib/agentic-search.ts (trimmed)",
)

para(
    "**What actually happens inside this one call:** Claude reads the CV and the user's target role. "
    "It decides which Apify actor to run (say, the LinkedIn scraper), what query to send it, and what country/location to set. "
    "Anthropic forwards the call to Apify MCP, gets the results back, feeds them back into Claude. "
    "Claude looks at what came back. If the results are weak, it tries a different query or a different actor. "
    "When it has a solid set, it calls our `finalize_jobs` tool with the final list. "
    "From our code's perspective this is **one** API call — all the tool round-trips happen inside it."
)

para(
    "**Why the custom `finalize_jobs` tool?** We need structured JSON at the end, not a blob of text to parse. "
    "A tool call with a typed schema guarantees the shape. No fragile regex, no 'Claude forgot a comma' bugs."
)

h3("Step 5 — The results page polls")
para(
    "Location: `components/search-poller.tsx`. "
    "The page loaded right after the form was submitted. It calls `GET /api/search?id=...` every couple of seconds. "
    "When the row flips from `running` to `complete`, it stops polling and shows the jobs."
)

# ============================================================
# WRAP UP
# ============================================================
h1("Summary — the whole thing in 30 seconds")
bullet("**Next.js layout:** `page.tsx` = a page, `route.ts` = an API endpoint, `'use client'` = runs in the browser, `middleware` = gatekeeper.")
bullet("**Login:** Supabase magic link. We just redirect back from the email link, exchange the code, and ensure a profile row exists.")
bullet("**CV upload:** form sends the PDF → server extracts text with `unpdf`, stores the file in Supabase Storage, asks Claude to structure it, saves it as the active CV.")
bullet(
    "**Search:** API route returns instantly with a search ID; a background job runs cheap sources and an **agentic** Claude-driven Apify loop in parallel; results are scored and written to the DB; the UI polls until done."
)
bullet(
    "**The agentic trick:** we attach Apify's MCP server to a single Claude API call. Claude picks queries, inspects results, refines, and returns a curated list via a typed `finalize_jobs` tool. Same mechanism that makes Claude Desktop's tool use feel smart — dropped into our backend."
)

doc.save(OUT)
print(f"wrote {OUT}")
