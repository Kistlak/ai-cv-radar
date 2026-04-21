"""
Generates docs/supabase-and-database-in-nextjs.docx.
A beginner-friendly explainer covering:
  1. What Supabase is and how this project uses it.
  2. How a database works with Next.js and an ORM (Drizzle).
Written for someone new to Next.js.
One-shot script - safe to delete after use.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "docs" / "supabase-and-database-in-nextjs.docx"

doc = Document()

for style_name in ("Normal",):
    s = doc.styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)


def shade(paragraph, color_hex="F4F4F4"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _inline(p, text):
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                p.add_run(text[i:])
                break
            r = p.add_run(text[i + 2:end])
            r.bold = True
            i = end + 2
        elif text[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                p.add_run(text[i:])
                break
            r = p.add_run(text[i + 1:end])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            i = end + 1
        else:
            nxt = len(text)
            for marker in ("**", "`"):
                pos = text.find(marker, i)
                if pos != -1 and pos < nxt:
                    nxt = pos
            p.add_run(text[i:nxt])
            i = nxt


def para(text):
    p = doc.add_paragraph()
    _inline(p, text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    _inline(p, text)


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    _inline(p, text)


def code(snippet, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for line in snippet.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        shade(p, "F4F4F4")
    doc.add_paragraph()


def callout(text, label="TIP"):
    p = doc.add_paragraph()
    shade(p, "FFF8E1")
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x57, 0x00)
    _inline(p, text)


# ============================================================
# TITLE
# ============================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Supabase and the Database in Next.js")
r.bold = True
r.font.size = Pt(22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("A beginner-friendly tour of how AI CV Radar stores data, authenticates users, and talks to Postgres")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# ============================================================
# PART 1 - WHAT IS SUPABASE
# ============================================================
h1("Part 1. What is Supabase?")

para(
    "Think of Supabase as a bundle of services you would otherwise have to build or rent separately, wrapped around a Postgres database. "
    "You sign up, create a project, and you instantly get four things:"
)

bullet("**A Postgres database** (SQL, same as you would run locally, but hosted for you).")
bullet("**Authentication** (sign-up, sign-in, password reset, OAuth - all handled).")
bullet("**Storage** (like a tiny S3 for files - we use it for uploaded CV PDFs).")
bullet("**Row Level Security (RLS)** - per-row access rules written in SQL, enforced at the database.")

para(
    "Everything is accessible via a REST API, a JavaScript client, and a standard Postgres connection string. "
    "That matters because it means we can pick-and-choose: use Supabase's auth and storage from the browser, "
    "but talk to the database using plain Postgres tools (Drizzle, in our case)."
)

callout(
    "You do not have to use all of Supabase. Many projects use Supabase purely for auth, or purely for its Postgres hosting. "
    "We happen to use auth + storage + database - three of the four."
)

# ============================================================
# PART 1b - HOW THIS PROJECT USES IT
# ============================================================
h1("How this project uses Supabase")

h2("1. Authentication")
para(
    "When a user signs up or logs in, Supabase handles it. "
    "Our own code never sees the password. Supabase sets an encrypted cookie in the browser, "
    "and on every request we ask Supabase \"who is this cookie?\" to get the current user."
)
code(
    "// lib/supabase/server.ts - used inside server components and API routes\n"
    "import { createServerClient } from '@supabase/ssr'\n"
    "import { cookies } from 'next/headers'\n"
    "\n"
    "export async function createClient() {\n"
    "  const cookieStore = await cookies()\n"
    "  return createServerClient(\n"
    "    process.env.NEXT_PUBLIC_SUPABASE_URL!,\n"
    "    process.env.NEXT_PUBLIC_SUPABASE_ANON_KEY!,\n"
    "    { cookies: { /* read/write session cookie */ } }\n"
    "  )\n"
    "}",
    caption="lib/supabase/server.ts",
)
para(
    "Anywhere we need the logged-in user, we do this:"
)
code(
    "const supabase = await createClient()\n"
    "const { data: { user } } = await supabase.auth.getUser()\n"
    "if (!user) return NextResponse.json({ error: 'Unauthorized' }, { status: 401 })",
    caption="typical auth check in an API route",
)

h2("2. Storage (uploaded CV PDFs)")
para(
    "When a user uploads a CV, we save the PDF itself to a **Supabase Storage bucket called `cvs`**. "
    "Think of a bucket as a folder in the cloud with access rules. The bucket is private - nobody can read a file without a signed URL or a valid session."
)
code(
    "// app/api/cv/upload/route.ts\n"
    "const filePath = `${user.id}/${Date.now()}.pdf`\n"
    "await supabase.storage\n"
    "  .from('cvs')\n"
    "  .upload(filePath, buffer, { contentType: 'application/pdf', upsert: true })",
    caption="uploading the PDF",
)
para(
    "Storing only the **path** in the database (`file_path` column on the `cvs` table) keeps the database small - "
    "the binary itself lives in Storage, and we fetch it only when needed."
)

h2("3. Postgres - the actual database")
para(
    "This is where everything else lives: user profiles, uploaded CVs, searches, job results, cached AI outputs. "
    "You access it two ways, for two different purposes:"
)
bullet(
    "**From Supabase's JS client** (what we use for auth, storage, and a few small reads) - "
    "goes through Supabase's REST API and respects Row Level Security."
)
bullet(
    "**From Drizzle ORM via a direct Postgres connection** (what we use for almost all data work) - "
    "a classic server-to-database connection with full SQL power, using the `DATABASE_URL` env var."
)
para(
    "Both are the same database. Just two different doors into it. Drizzle is more powerful; the Supabase client is more convenient for auth."
)

# ============================================================
# PART 2 - DATABASES WITH NEXT.JS
# ============================================================
h1("Part 2. How the database works with Next.js")

para(
    "Before diving into code, one thing you need to internalize about Next.js: **there are two worlds**, and the database lives only in one of them."
)

h2("Server vs Client")
bullet(
    "**Server components / API routes / server actions** - run on the server (like a traditional Node.js backend). "
    "These can read env vars, open database connections, and return HTML or JSON."
)
bullet(
    "**Client components** (`'use client'` at the top) - run in the browser. "
    "These can use React state, `useEffect`, and the DOM. They **cannot** touch the database directly - they have to ask the server."
)

para(
    "This split matters because a database connection string (`DATABASE_URL`) is a secret. "
    "If you imported the database module in a client component, that secret would get bundled into the JavaScript the browser downloads - "
    "which means anyone on the internet could read it. So Next.js is strict: database code lives server-side only."
)

callout(
    "Rule of thumb: if a file has `'use client'` at the top, it must never import `@/db`. "
    "Instead, it makes `fetch('/api/...')` calls to hit server-side routes that do the DB work."
)

h2("Where server-side data lives in this project")
para(
    "There are three places where we run database code, and you will see this pattern in every Next.js app:"
)
numbered(
    "**Server components** (files like `app/(app)/cv/page.tsx`) - these are `async` functions. "
    "They query the database at render time and return ready-made HTML. Fastest path for read-only pages."
)
numbered(
    "**API routes** (files like `app/api/cv/upload/route.ts`) - classic endpoints. "
    "Exported `POST`, `GET`, `DELETE` functions. Used for mutations, uploads, and anything a client component needs to call."
)
numbered(
    "**Server-side helpers** (`lib/run-search.ts`, `lib/job-ai-helpers.ts`) - shared code imported by the above. "
    "These never run in the browser."
)

# ============================================================
# PART 3 - WHAT IS AN ORM
# ============================================================
h1("Part 3. What is an ORM, and why Drizzle?")

para(
    "**ORM** stands for Object-Relational Mapper. In plain English: a library that lets you write database queries "
    "in your programming language (TypeScript, in our case) instead of raw SQL strings."
)

para("Without an ORM:")
code(
    "const result = await client.query(\n"
    "  'SELECT id, email FROM profiles WHERE id = $1',\n"
    "  [userId]\n"
    ")\n"
    "// result.rows is typed as `any`. Typos in column names blow up at runtime.",
    caption="raw SQL - works, but fragile",
)

para("With Drizzle:")
code(
    "const [profile] = await db\n"
    "  .select({ id: profiles.id, email: profiles.email })\n"
    "  .from(profiles)\n"
    "  .where(eq(profiles.id, userId))\n"
    "  .limit(1)\n"
    "// `profile` is fully typed. Typo `profiles.emial`? TypeScript catches it before you run the code.",
    caption="same query in Drizzle",
)

para(
    "The two big wins:"
)
bullet("**Type safety** - the shape of every row is known at compile time. Renames, typos, missing columns become red squiggles in your editor.")
bullet("**Composable queries** - you can build queries up programmatically without string-concatenating SQL (and without SQL injection risks).")

h2("Why Drizzle specifically?")
para(
    "Drizzle is lightweight and stays close to SQL - if you know SQL, you can read Drizzle code in a minute. "
    "It does not try to hide the database. It also integrates cleanly with Supabase's Postgres and with Next.js server runtimes."
)

# ============================================================
# PART 4 - THE SCHEMA
# ============================================================
h1("Part 4. The schema - your tables as TypeScript")

para(
    "In Drizzle, you do not write `CREATE TABLE` statements by hand. "
    "You describe your tables in a TypeScript file (`db/schema.ts`) and Drizzle can generate (or push) the SQL for you."
)

code(
    "// db/schema.ts - simplified excerpt\n"
    "export const profiles = pgTable('profiles', {\n"
    "  id: uuid('id').primaryKey(),\n"
    "  email: text('email').notNull(),\n"
    "  createdAt: timestamp('created_at', { withTimezone: true }).defaultNow().notNull(),\n"
    "})\n"
    "\n"
    "export const cvs = pgTable('cvs', {\n"
    "  id: uuid('id').primaryKey().defaultRandom(),\n"
    "  userId: uuid('user_id')\n"
    "    .notNull()\n"
    "    .references(() => profiles.id, { onDelete: 'cascade' }),\n"
    "  filePath: text('file_path').notNull(),\n"
    "  rawText: text('raw_text').notNull(),\n"
    "  structured: jsonb('structured').notNull(),\n"
    "  generalCv: jsonb('general_cv'),\n"
    "  generalCoverLetter: text('general_cover_letter'),\n"
    "  isActive: boolean('is_active').default(true).notNull(),\n"
    "})",
    caption="db/schema.ts",
)

para(
    "A few things to notice:"
)
bullet("**Column types are explicit** - `uuid`, `text`, `jsonb`, `timestamp`, `boolean`. These map 1:1 to Postgres types.")
bullet("**Foreign keys are readable** - `.references(() => profiles.id, { onDelete: 'cascade' })` means \"this row points at a profile, and if that profile is deleted, cascade-delete this too\".")
bullet("**`notNull()` and `defaultNow()`** are the same constraints you would write in SQL, just named nicer.")

# ============================================================
# PART 5 - THE DB CLIENT
# ============================================================
h1("Part 5. The db client - one connection, shared everywhere")

code(
    "// db/index.ts\n"
    "import { drizzle } from 'drizzle-orm/postgres-js'\n"
    "import postgres from 'postgres'\n"
    "import * as schema from './schema'\n"
    "\n"
    "const connectionString = process.env.DATABASE_URL!\n"
    "\n"
    "const client = postgres(connectionString, { prepare: false })\n"
    "export const db = drizzle(client, { schema })",
    caption="db/index.ts - the whole file",
)

para(
    "That is it. Three lines of real code. `db` is now an object you can import anywhere server-side:"
)
code(
    "import { db } from '@/db'\n"
    "import { cvs } from '@/db/schema'\n"
    "import { eq } from 'drizzle-orm'\n"
    "\n"
    "const [cv] = await db\n"
    "  .select()\n"
    "  .from(cvs)\n"
    "  .where(eq(cvs.userId, user.id))\n"
    "  .limit(1)",
    caption="reading a CV for the logged-in user",
)

callout(
    "The `@/` prefix is a TypeScript alias configured in `tsconfig.json`. It means \"the project root\". "
    "So `@/db` resolves to `./db/index.ts`. Next.js and TypeScript agree on this mapping."
)

# ============================================================
# PART 6 - A FULL ROUND TRIP
# ============================================================
h1("Part 6. A full round trip - clicking \"Upload CV\"")

para(
    "Let us trace one user action end to end. The user drops a PDF into the upload area on the CV page. Here is what happens, step by step:"
)

numbered(
    "**Browser (client component):** `components/cv-upload-form.tsx` reads the file from the drag-drop event "
    "and calls `fetch('/api/cv/upload', { method: 'POST', body: formData })`."
)
numbered(
    "**Server (API route):** `app/api/cv/upload/route.ts` starts running. "
    "First it calls `supabase.auth.getUser()` to confirm the user is logged in - the session cookie does this job."
)
numbered(
    "**Decrypt API keys:** Next, it looks up the user's Anthropic key from the `user_api_keys` table, decrypts it "
    "(we store it encrypted with `APP_ENCRYPTION_KEY`), and keeps it in memory for this one request."
)
numbered(
    "**Parse the PDF:** `unpdf` extracts the raw text from the PDF bytes."
)
numbered(
    "**Save the file to Storage:** the PDF is uploaded to the `cvs` bucket under `{userId}/{timestamp}.pdf`."
)
numbered(
    "**Call Claude:** Claude receives the raw text and returns a JSON object: name, email, skills, experience, education."
)
numbered(
    "**Deactivate old CV + insert new one:** using Drizzle, we mark the user's existing CV row `isActive = false`, then `db.insert(cvs).values({...}).returning()` to save the new one."
)
numbered(
    "**Respond:** the API route returns `{ cv: newCv }` as JSON. The browser gets a 200."
)
numbered(
    "**Re-render:** the client calls `router.refresh()`, which tells Next.js to re-run the server component for the CV page. "
    "The server component re-queries the DB and returns new HTML with the parsed data visible."
)

callout(
    "Notice how the **client does no data work**. It hands off to the server, then asks for a re-render. "
    "That is the Next.js App Router pattern in one sentence."
)

# ============================================================
# PART 7 - DEVELOPMENT WORKFLOW
# ============================================================
h1("Part 7. Day-to-day developer workflow")

h2("Adding a column or a new table")
numbered("Edit `db/schema.ts` - add the column or new `pgTable(...)` block.")
numbered("Run `npx drizzle-kit push` - Drizzle diffs your schema against the live Supabase database and applies the change.")
numbered("Or, for production discipline, run `npx drizzle-kit generate` to produce a SQL migration file in `supabase/migrations/`, review it, then apply it with `drizzle-kit migrate` or via the Supabase SQL editor.")

para(
    "In this project you will see both patterns: hand-written migration SQL files in `supabase/migrations/` "
    "(for clarity and review), alongside the schema which is the source of truth in TypeScript."
)

h2("Reading and writing data")
para("A cheat sheet of the most common Drizzle patterns you will hit:")

code(
    "// SELECT\n"
    "const rows = await db.select().from(cvs).where(eq(cvs.userId, userId))\n"
    "\n"
    "// SELECT with specific columns and a join\n"
    "const rows = await db\n"
    "  .select({ job: jobResults, search: searches })\n"
    "  .from(jobResults)\n"
    "  .innerJoin(searches, eq(jobResults.searchId, searches.id))\n"
    "  .where(and(eq(jobResults.id, jobId), eq(searches.userId, userId)))\n"
    "\n"
    "// INSERT\n"
    "await db.insert(cvs).values({ userId, filePath, rawText, structured })\n"
    "\n"
    "// INSERT and get the row back\n"
    "const [newCv] = await db.insert(cvs).values({...}).returning()\n"
    "\n"
    "// UPDATE\n"
    "await db.update(cvs).set({ isActive: false }).where(eq(cvs.id, cvId))\n"
    "\n"
    "// DELETE\n"
    "await db.delete(cvs).where(eq(cvs.id, cvId))",
    caption="the Drizzle patterns you will use 95% of the time",
)

# ============================================================
# PART 8 - GOTCHAS
# ============================================================
h1("Part 8. Gotchas a Next.js newcomer will hit")

h2("\"Cannot find module 'net'\"")
para(
    "You probably imported `@/db` inside a client component. Move the import into a server component or an API route, "
    "and call that route from the client with `fetch(...)`."
)

h2("\"user is null\" in an API route")
para(
    "Usually means the session cookie is not being forwarded. Double-check you are calling `createClient()` from `lib/supabase/server.ts` "
    "(not the browser client), and that the user is actually logged in."
)

h2("Schema drift")
para(
    "If you edited `db/schema.ts` but forgot to `drizzle-kit push`, TypeScript will still compile, but queries will blow up at runtime "
    "(\"column does not exist\"). When in doubt, push."
)

h2("Row Level Security blocking a query")
para(
    "Supabase tables can have RLS policies. When we use the **service role** connection (via `DATABASE_URL` + Drizzle), RLS is bypassed - "
    "that is intentional, because the server is trusted. When we use the **anon** Supabase client from the browser, RLS is enforced. "
    "If you see \"permission denied for table ...\" from the browser, you need an RLS policy."
)

# ============================================================
# SUMMARY
# ============================================================
h1("The whole thing in 30 seconds")

bullet("**Supabase** is Postgres + auth + storage, all hosted.")
bullet("We use Supabase's **auth** for login (cookies), **storage** for CV PDFs, and its **Postgres** for everything else.")
bullet("We talk to Postgres through **Drizzle**, a TypeScript ORM. The schema lives in `db/schema.ts`; the connection is in `db/index.ts`.")
bullet("Next.js splits code into **server** (DB, env vars, API routes) and **client** (browser, React state). Database code only ever runs on the server.")
bullet("The typical loop: **client fetch -> server API route -> Drizzle query -> JSON response -> router.refresh**.")

doc.save(OUT)
print(f"wrote {OUT}")
