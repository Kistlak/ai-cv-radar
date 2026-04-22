"""
Generates docs/nextjs-beginner-guide.docx.
A simple, plain-English explanation of what is in this project and why,
written for someone who is brand new to Next.js.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "docs" / "nextjs-beginner-guide.docx"

doc = Document()

# ---------- base style ----------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)


# ---------- helpers ----------
def shade(paragraph, color_hex="F4F4F4"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _inline(p, text):
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                p.add_run(text[i:])
                break
            run = p.add_run(text[i + 2:end])
            run.bold = True
            i = end + 2
        elif text[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                p.add_run(text[i:])
                break
            run = p.add_run(text[i + 1:end])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            i = end + 1
        else:
            nxt = len(text)
            for marker in ("**", "`"):
                pos = text.find(marker, i)
                if pos != -1 and pos < nxt:
                    nxt = pos
            p.add_run(text[i:nxt])
            i = nxt


def para(text):
    p = doc.add_paragraph()
    _inline(p, text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    _inline(p, text)


def code(snippet, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for line in snippet.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        shade(p)
    doc.add_paragraph()


# ============================================================
# TITLE
# ============================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("AI CV Radar — A Beginner's Guide")
r.bold = True
r.font.size = Pt(22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("What I built, and why — explained for someone new to Next.js")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()


# ============================================================
# 1. WHAT THE APP DOES
# ============================================================
h1("1. What this app does (in one paragraph)")
para(
    "AI CV Radar is a website where you upload your CV once, search for jobs across "
    "sites like LinkedIn, Indeed, Remotive and Adzuna, and the app uses AI (Claude) "
    "to score how well each job matches you. For any job you like, you can click one "
    "button to get a **deep-dive fit analysis**, a **tailored cover letter**, or a "
    "**Word-document CV rewritten for that role**. There is also a **Chrome extension** "
    "that autofills application forms with your details — you still click submit yourself."
)


# ============================================================
# 2. NEXT.JS IN 60 SECONDS
# ============================================================
h1("2. Next.js in 60 seconds (only the bits we use)")
para(
    "Next.js is a framework built on top of React. You can think of it as "
    "**React + a backend + routing + build tools**, all in one project. "
    "Almost everything you need to know to read this codebase comes down to a few "
    "file-name rules."
)

bullet("`page.tsx` — becomes a page you can visit in the browser. The URL mirrors the folder path. `app/settings/page.tsx` → `/settings`.")
bullet("`layout.tsx` — wraps every page inside its folder (shared header, theme, auth-check, etc.).")
bullet("`route.ts` — a backend API endpoint. No Express needed. `app/api/cv/upload/route.ts` → `POST /api/cv/upload`.")
bullet("**Server Components (the default).** The file runs on the server. It can read the database directly, use secret keys, and returns ready-made HTML to the browser. Fast and safe.")
bullet("**Client Components (`'use client'` at the top).** The file runs in the browser. Needed whenever you use `useState`, `onClick`, form interactions, or anything that moves on the page.")
bullet("**Folders in `(parentheses)`** group files without adding anything to the URL. We use `(app)` for logged-in pages and `(auth)` for sign-in / sign-up.")
bullet("**`[brackets]`** in a folder name mean a dynamic URL parameter. `search/[id]/page.tsx` → `/search/abc123` and the page receives `id = \"abc123\"`.")

para("That's genuinely enough to navigate the whole project.")


# ============================================================
# 3. TECH STACK
# ============================================================
h1("3. The tech stack and why each piece exists")

h2("Next.js 16 + React 19 + TypeScript")
para(
    "**Next.js** gives us pages, routing, and a backend in one project — so I don't have to "
    "run a separate server. **React** is how we build the UI out of components. "
    "**TypeScript** adds types to JavaScript, which catches bugs before the page even loads."
)

h2("Tailwind CSS 4 + Base UI + shadcn-style components")
para(
    "**Tailwind** lets me style things by writing class names (`px-4 text-sm rounded-md`) "
    "instead of writing separate CSS files. It's faster and the styles never drift out of "
    "sync with the markup. **Base UI** gives me accessible building blocks (dialogs, "
    "dropdowns, tooltips) that I style with Tailwind. Lives in `components/ui/`."
)

h2("Supabase (database + auth + file storage)")
para(
    "**Supabase** is a hosted Postgres database with a bunch of extras baked in. "
    "I use three parts of it:"
)
bullet("**Auth** — sign up, sign in, sessions (stored as cookies, checked server-side).")
bullet("**Postgres** — where all CVs, searches, and job results live.")
bullet("**Storage** — a private bucket called `cvs` that holds the uploaded PDF files themselves.")
para(
    "Supabase basically saves me from writing a whole user system from scratch. "
    "See `lib/supabase/server.ts` and `lib/supabase/client.ts`."
)

h2("Drizzle ORM")
para(
    "An **ORM** (Object-Relational Mapper) is a library that lets me work with the "
    "database in TypeScript instead of writing raw SQL strings. **Drizzle** is one of "
    "the lightweight ones — it stays close to SQL, so I can learn one and get the other."
)
para("Two benefits:")
bullet("**Type-safety.** If I misspell a column name, the code won't compile.")
bullet("**Migrations.** I edit `db/schema.ts`, run `npx drizzle-kit push`, and Drizzle updates the actual database tables for me.")
code(
    "// db/schema.ts — the shape of a single table\n"
    "export const cvs = pgTable('cvs', {\n"
    "  id: uuid('id').primaryKey().defaultRandom(),\n"
    "  userId: uuid('user_id').notNull().references(() => profiles.id),\n"
    "  rawText: text('raw_text').notNull(),\n"
    "  structured: jsonb('structured').notNull(),\n"
    "  isActive: boolean('is_active').default(true).notNull(),\n"
    "})",
    caption="A Drizzle table definition. No SQL needed.",
)
code(
    "// Using it in an API route\n"
    "const rows = await db.select().from(cvs).where(eq(cvs.userId, user.id))",
    caption="Type-safe query — Drizzle knows `rows[0].rawText` is a string.",
)

h2("Anthropic Claude (the AI)")
para(
    "All AI work goes through the **Claude API** via the official "
    "`@anthropic-ai/sdk` package. **Sonnet 4.6** is used for generation (cover letters, "
    "tailored CVs, deep dives). **Haiku 4.5** — a smaller, faster model — is used for "
    "scoring jobs in bulk because it's roughly 5× cheaper per call."
)

h2("docx and unpdf (document work)")
bullet("**unpdf** — pulls plain text out of a user's uploaded PDF CV.")
bullet("**docx** — builds a real Microsoft Word `.docx` file in JavaScript, so the tailored CV you download opens cleanly in Word.")

h2("Playwright (end-to-end tests)")
para(
    "**Playwright** is a testing tool that actually launches a real Chromium browser, "
    "clicks buttons on my site, and checks things work. Files live in `tests/e2e/`."
)


# ============================================================
# 4. FOLDER LAYOUT
# ============================================================
h1("4. The folder layout — what lives where")
code(
    "ai-cv-radar/\n"
    "├── app/                  Every page and API route lives here\n"
    "│   ├── (app)/            Logged-in pages (dashboard, cv, search, settings)\n"
    "│   ├── (auth)/           Sign in / sign up pages\n"
    "│   ├── api/              Backend endpoints (route.ts files)\n"
    "│   ├── layout.tsx        Root layout — fonts, theme, <html> tag\n"
    "│   ├── page.tsx          The public landing page at /\n"
    "│   └── globals.css       Tailwind imports + a few base styles\n"
    "│\n"
    "├── components/           Reusable React components (forms, dialogs, nav)\n"
    "│   └── ui/               Base-UI / shadcn primitives (Button, Dialog, ...)\n"
    "│\n"
    "├── db/                   Drizzle ORM\n"
    "│   ├── schema.ts         Table definitions\n"
    "│   └── index.ts          The `db` client you import everywhere\n"
    "│\n"
    "├── lib/                  Non-UI code — the real logic\n"
    "│   ├── supabase/         Auth + storage helpers (server / client / middleware)\n"
    "│   ├── job-sources/      One file per job board (remotive, adzuna, ...)\n"
    "│   ├── run-search.ts     Orchestrates one full search end-to-end\n"
    "│   ├── agentic-search.ts Claude-driven search loop (uses Apify MCP)\n"
    "│   ├── score-jobs.ts     Scores jobs 0-100 against the user's CV\n"
    "│   ├── cv-docx.ts        Builds the downloadable Word document\n"
    "│   └── crypto.ts         Encrypts user API keys before saving\n"
    "│\n"
    "├── extension/            The Chrome/Edge browser extension (autofill)\n"
    "├── tests/e2e/            Playwright tests\n"
    "├── supabase/migrations/  SQL migrations generated by Drizzle\n"
    "└── scripts/              One-off utility scripts (like this one)",
)

para(
    "The split is simple: **`app/` is routing and pages**, **`components/` is UI pieces**, "
    "**`db/` is database shape**, **`lib/` is everything else** (API calls, AI, parsing, etc.). "
    "Each folder has one job."
)


# ============================================================
# 5. HOW A REQUEST FLOWS
# ============================================================
h1("5. How a typical user action flows through the code")

h2("Example: user uploads their CV")
bullet("The browser renders `components/cv-upload-form.tsx` — a **client component** because it has a file input and state.")
bullet("User picks a PDF and clicks upload. The form POSTs to `/api/cv/upload`.")
bullet("`app/api/cv/upload/route.ts` runs **on the server**. It checks the session with Supabase, saves the PDF to Supabase Storage, and extracts the text with `unpdf`.")
bullet("It sends that text to Claude (`@anthropic-ai/sdk`), which returns a structured JSON version of the CV (skills, experience, education).")
bullet("It stores everything in the `cvs` table using Drizzle.")
bullet("The browser gets a success response and navigates to `/cv`, which is a **server component** that reads the new row straight from the database and renders it.")

h2("Example: user runs a job search")
bullet("`components/search-form.tsx` collects query, location, and job sources.")
bullet("It POSTs to `/api/search`, which inserts a row in the `searches` table with status `running` and kicks off `lib/run-search.ts` in the background.")
bullet("`run-search.ts` calls each source adapter in `lib/job-sources/` (they all look the same: `fetchJobs(query) → Job[]`).")
bullet("Results go into the `job_results` table, then `lib/score-jobs.ts` asks Claude Haiku to score them 0-100.")
bullet("`components/search-poller.tsx` (client-side) polls `/api/search/[id]` every second and refreshes the page when the status flips to `done`.")


# ============================================================
# 6. DATABASE DESIGN
# ============================================================
h1("6. The database, in five tables")

para("Everything lives in Postgres (managed by Supabase). Here are the five tables:")

bullet("**`profiles`** — one row per user. Just `id`, `email`, `created_at`. Linked to Supabase Auth by matching `id`.")
bullet("**`user_api_keys`** — user-provided API keys (Anthropic, Apify, Adzuna, RapidAPI). **Encrypted at rest** using `lib/crypto.ts` — even if someone dumped the DB, the keys are useless without our `APP_ENCRYPTION_KEY`.")
bullet("**`cvs`** — uploaded CVs. Stores the raw extracted text, a structured JSON version, plus an optional polished general CV and cover letter.")
bullet("**`searches`** — one row per job search. Tracks status (`running` / `done` / `error`) so the UI can show progress and you can cancel.")
bullet("**`job_results`** — one row per job found. Holds the match score, match reason, and the cached AI outputs (deep dive, cover letter, tailored CV). Caching these means the user only pays for Claude once per job.")

para(
    "All tables use **`onDelete: 'cascade'`** on their foreign keys. That means if I "
    "delete a user, their CVs, searches, and job results are wiped too — no orphan data."
)

para("There's also one index for speed:")
code(
    "index('job_results_search_score_idx').on(\n"
    "  table.searchId, table.matchScore,\n"
    ")",
    caption="Lets us pull 'top 20 jobs from this search by score' in one quick lookup.",
)


# ============================================================
# 7. PRINCIPLES I STUCK TO
# ============================================================
h1("7. Principles I stuck to")

h2("Keep auth on the server")
para(
    "Every logged-in page lives under `app/(app)/`. The `layout.tsx` at the top of that "
    "folder calls `supabase.auth.getUser()` on the server, and if there is no user it "
    "redirects to `/login` before the page is ever rendered. I don't trust the browser "
    "to hide private content — the server decides."
)

h2("Never trust the client with secrets")
para(
    "Anthropic and Apify keys live in the database, encrypted. API routes look them up "
    "server-side right before calling Claude. The browser never sees them."
)

h2("Cache AI output")
para(
    "Claude calls cost real money. Every AI result (deep dive, cover letter, tailored CV) "
    "is stored on the job row, so re-opening a job is free. A **Regenerate** button is "
    "the only thing that actually re-hits Claude."
)

h2("One job per file")
para(
    "Each job source (Remotive, Adzuna, LinkedIn, ...) is its own file in `lib/job-sources/` "
    "with the same shape: `fetchJobs(query) → Job[]`. Adding a new source is one file, not a refactor."
)

h2("Types flow end-to-end")
para(
    "Drizzle infers TypeScript types from the schema (`typeof cvs.$inferSelect`), so a "
    "CV row has the exact same type whether I'm reading it in an API route, passing it "
    "to a React component, or sending it to Claude. Rename a column → TypeScript tells "
    "me every place I need to fix."
)


# ============================================================
# 8. ARCHITECTURE AT A GLANCE
# ============================================================
h1("8. Architecture at a glance")
code(
    "   Browser                Next.js server                 External\n"
    "   -------                --------------                 --------\n"
    "  [page.tsx] ── fetch ──► [route.ts]  ── Drizzle ──►   Postgres (Supabase)\n"
    "  [components]            [lib/*]     ── SDK call ──►  Claude API\n"
    "                                      ── HTTP    ──►  Remotive / Adzuna / Apify\n"
    "                                      ── upload  ──►  Supabase Storage (PDFs)\n"
    "\n"
    "  [Chrome ext.] ── fetch ──► /api/me/application-profile (with cookie auth)",
)

para(
    "The server is the only thing that talks to Claude, Supabase, and the job boards. "
    "The browser only talks to the server. That one rule keeps secrets safe and keeps "
    "the code tidy."
)


# ============================================================
# 9. HANDY COMMANDS
# ============================================================
h1("9. Commands you'll actually use")
code(
    "npm run dev             # start the dev server at localhost:3000\n"
    "npm run build           # production build (also type-checks)\n"
    "npm run lint            # run ESLint\n"
    "npm run test:e2e        # run Playwright browser tests\n"
    "npx drizzle-kit push    # apply db/schema.ts changes to Supabase",
)


# ============================================================
# 10. IF YOU WANT TO LEARN MORE
# ============================================================
h1("10. If you want to learn more")
bullet("**Next.js App Router** — nextjs.org/docs/app. Read about pages, layouts, and route handlers. That's ~80% of this project.")
bullet("**React Server Components vs Client Components** — the single biggest mental-model shift. Worth one good read.")
bullet("**Drizzle** — orm.drizzle.team. Look at the Postgres quickstart.")
bullet("**Supabase Auth with Next.js** — supabase.com/docs/guides/auth/server-side/nextjs. Explains the cookie-based session dance.")
bullet("**Tailwind** — tailwindcss.com/docs. Skim the utilities, don't memorize them. You look things up as you go.")

para(
    "That's the whole tour. Open `app/(app)/search/page.tsx` and follow the imports — "
    "every concept above shows up within two or three files."
)

# ============================================================
# SAVE
# ============================================================
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"Wrote {OUT}")
