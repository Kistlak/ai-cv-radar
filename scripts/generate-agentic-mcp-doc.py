"""
Generates docs/agentic-search-with-apify-mcp.docx.
A friendly, plain-English explainer for non-technical readers.
One-shot script — safe to delete after use.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = Path(__file__).resolve().parent.parent / "docs" / "agentic-search-with-apify-mcp.docx"

doc = Document()

for style_name in ("Normal",):
    s = doc.styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(11)

section = doc.sections[0]
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)


def shade(paragraph, color_hex="F4F4F4"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _inline(p, text):
    i = 0
    while i < len(text):
        if text.startswith("**", i):
            end = text.find("**", i + 2)
            if end == -1:
                p.add_run(text[i:])
                break
            r = p.add_run(text[i + 2 : end])
            r.bold = True
            i = end + 2
        elif text[i] == "`":
            end = text.find("`", i + 1)
            if end == -1:
                p.add_run(text[i:])
                break
            r = p.add_run(text[i + 1 : end])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            i = end + 1
        else:
            nxt = len(text)
            for marker in ("**", "`"):
                pos = text.find(marker, i)
                if pos != -1 and pos < nxt:
                    nxt = pos
            p.add_run(text[i:nxt])
            i = nxt


def para(text):
    p = doc.add_paragraph()
    _inline(p, text)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    _inline(p, text)


def code(snippet, caption=None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    for line in snippet.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.1)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        shade(p, "F4F4F4")
    doc.add_paragraph()


def callout(text):
    p = doc.add_paragraph()
    shade(p, "FFF8E1")
    r = p.add_run("TIP: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x8B, 0x57, 0x00)
    _inline(p, text)


# ============================================================
# TITLE
# ============================================================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Agentic Job Search")
r.bold = True
r.font.size = Pt(22)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("How AI CV Radar uses Apify MCP and Claude to find jobs that actually fit you")
r.italic = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()

# ============================================================
# SECTION 1 — Plain-English intros
# ============================================================
h1("First, three words you will keep seeing")

h2("Claude")
para(
    "Claude is an AI assistant built by Anthropic — the same family of models behind ChatGPT-style chat apps. "
    "In this project we call Claude from our own backend through a developer API, so the app can ask Claude questions "
    "without a person being in the loop. Think of it as having a very fast research assistant on tap."
)

h2("Apify")
para(
    "Apify is a cloud service that runs ready-made scrapers for sites that do not offer good APIs. "
    "Want 50 recent LinkedIn job postings for \"React developer in Berlin\"? There is an Apify actor for that. "
    "Indeed, Glassdoor — same story. You pay per run, and the results come back as clean JSON."
)

h2("MCP (Model Context Protocol)")
para(
    "MCP is a recent standard that lets AI models use external tools in a predictable way. "
    "It is the same plumbing that lets Claude Desktop search your files, browse the web, or run Apify scrapers for you — "
    "except now any app can plug into it, not just desktop chat clients. "
    "Apify hosts their own MCP server at `https://mcp.apify.com`, which exposes all their scrapers as tools."
)

callout(
    "If you have ever used Claude Desktop and watched it run a tool, you have already used MCP without knowing the name. "
    "We are taking that same mechanism and embedding it inside our backend."
)

# ============================================================
# SECTION 2 — Why agentic
# ============================================================
h1("The problem: dumb search gives dumb results")

para(
    "A traditional search pipeline looks like this: user types a query → we send that query to a job board → we display whatever comes back. "
    "It is fast and easy, but it breaks in obvious ways:"
)
bullet("If the user's query is slightly off (\"react dev\" vs \"React Engineer\"), the results are mostly junk.")
bullet("Different boards need different phrasings — one query does not fit all of them.")
bullet("If a query returns 3 bad matches, our code cannot notice and try again. It just shows the 3 bad matches.")

para(
    "When you run the exact same searches manually in Claude Desktop using Apify's tools, the output is noticeably better. "
    "Why? Because Claude iterates. It tries a query, skims the results, notices the wrong roles are showing up, "
    "refines the wording, tries a second actor, and stops only when the list looks strong. "
    "A human does the same thing without thinking about it. A hard-coded pipeline cannot."
)

para(
    "**The fix is to put Claude in the driver's seat.** Instead of our code calling Apify, we let Claude call Apify — "
    "picking the query, inspecting what comes back, refining, and handing us a curated list at the end. "
    "That is what \"agentic\" means in this project."
)

# ============================================================
# SECTION 3 — The workflow, visualized
# ============================================================
h1("The workflow, step by step")

para("Here is what happens after you click \"Search\":")

bullet("**1. You submit the form.** The browser sends your query, location, and preferences to our backend.")
bullet(
    "**2. We create a search row and return immediately.** We do not make you wait on the spinning wheel for a minute. "
    "The backend saves a `searches` row with status `running`, returns its ID, and launches the real work in the background."
)
bullet(
    "**3. Two paths run in parallel.** "
    "**(a)** Cheap, free-tier sources (Remotive, Adzuna, JSearch) — fast but often low relevance. "
    "**(b)** The agentic path — Claude driving Apify's LinkedIn / Indeed / Glassdoor scrapers. "
    "Both finish, we merge the results, remove duplicates."
)
bullet(
    "**4. A second Claude pass scores every job against your CV** (match score out of 100 + a one-sentence reason). "
    "We do this even after the agent has curated — two independent signals is better than one."
)
bullet(
    "**5. The results page, which your browser went to as soon as step 2 returned, polls every 3 seconds.** "
    "The moment the row flips to `complete`, you see the ranked list."
)

callout(
    "From the user's perspective: click search → see a loading screen → 60-90 seconds later → ranked jobs. "
    "From our code's perspective: one quick API call plus one background job with a lot going on inside it."
)

# ============================================================
# SECTION 4 — The agentic engine
# ============================================================
h1("Inside the agentic engine (the interesting bit)")

para(
    "The whole agent fits in one file: `lib/agentic-search.ts`. "
    "And within that file, the whole thing is really just one API call to Claude — with two things attached. "
    "Understanding those two attachments is the whole story."
)

h2("Attachment #1: the Apify MCP server")
para(
    "We tell Claude: \"you can talk to Apify's MCP server at `https://mcp.apify.com`, here is the user's Apify token\". "
    "That alone gives Claude access to every Apify scraper as a callable tool — LinkedIn, Indeed, Glassdoor, and hundreds more."
)

h2("Attachment #2: our own `finalize_jobs` tool")
para(
    "We also give Claude one custom tool that only our app knows about, called `finalize_jobs`. "
    "It has a strict schema: title, company, apply URL, location, and so on. "
    "We tell Claude: \"when you are done, call this exactly once with your final curated list — and never invent jobs\"."
)
para(
    "**Why a tool instead of just asking for JSON in the reply?** "
    "Tool calls come with a guaranteed shape. No fragile regex, no \"Claude forgot a comma\" bugs. "
    "Our code reads the structured payload directly."
)

h2("The call itself")
code(
    "const response = await client.beta.messages.create({\n"
    "  model: 'claude-sonnet-4-6',\n"
    "  system: buildSystemPrompt(targetCount),\n"
    "  messages: [{ role: 'user', content: buildUserPrompt(input) }],\n"
    "\n"
    "  // Attachment #1 — Apify's MCP server, with the user's own Apify token\n"
    "  mcp_servers: [{\n"
    "    type: 'url',\n"
    "    url: 'https://mcp.apify.com',\n"
    "    name: 'apify',\n"
    "    authorization_token: input.apifyToken,\n"
    "  }],\n"
    "\n"
    "  // Attachment #2 — all Apify tools, plus our local finalize_jobs\n"
    "  tools: [\n"
    "    { type: 'mcp_toolset', mcp_server_name: 'apify' },\n"
    "    FINALIZE_JOBS_TOOL,\n"
    "  ],\n"
    "\n"
    "  betas: ['mcp-client-2025-11-20'],\n"
    "})",
    caption="lib/agentic-search.ts — the one call that does everything",
)

h2("What actually happens inside that one call")
para(
    "From our code's perspective this is a single API call. From Claude's perspective it is a conversation that plays out server-side, completely hidden from us:"
)
bullet("Claude reads the CV and the user's target role.")
bullet("It picks an Apify actor — say, the LinkedIn scraper — and writes a tight query.")
bullet("Anthropic forwards that tool call to Apify's MCP server. Apify runs the scraper. Results come back.")
bullet("Claude looks at the results. If they are weak (wrong stack, wrong seniority, stale postings), it tries a different query or a different actor.")
bullet("Typically it runs 2-3 searches across actors before it is satisfied.")
bullet("Finally, it calls `finalize_jobs` with its curated list. That is the signal for us to read the result and move on.")

para(
    "**We do not write a tool-use loop on our side.** Anthropic handles the round-trips server-side. "
    "In one of the test runs we counted 12 Apify tool calls inside a single API request — all invisible to our code."
)

# ============================================================
# SECTION 5 — The system prompt
# ============================================================
h1("Telling Claude how to behave")

para(
    "The system prompt is how we shape Claude's behavior. Ours is short and opinionated. Here is the gist:"
)

code(
    "You are an agentic job-search assistant.\n"
    "Goal: return 15-30 genuinely relevant, recent job postings that fit this\n"
    "specific candidate's CV and target role.\n"
    "\n"
    "How to work:\n"
    "1. Read the CV and the candidate's target role.\n"
    "2. Pick ONE tight query and call an Apify actor.\n"
    "3. Inspect what came back. If results are off-target, refine and try again.\n"
    "4. Prefer 2-3 solid searches across different actors over one broad one.\n"
    "5. When you have a strong set, call finalize_jobs exactly once.\n"
    "\n"
    "Hard rules:\n"
    "- Never invent jobs. Every item must come from a tool result you actually saw.\n"
    "- Do not pad with irrelevant jobs to hit the count. 12 great jobs beat 30 mediocre ones.\n"
    "- Do not call finalize_jobs before running at least one Apify tool.\n"
    "- Do not ask the user clarifying questions — work with what you have.",
    caption="lib/agentic-search.ts — system prompt (slightly trimmed)",
)

para(
    "Three things worth pointing out:"
)
bullet(
    "**\"Never invent jobs.\"** Without this rule, an AI under pressure to return N items will sometimes make up postings. "
    "Pinning every item to a real tool result is non-negotiable."
)
bullet(
    "**\"12 great jobs beat 30 mediocre ones.\"** The model naturally wants to hit the requested count. "
    "We explicitly tell it: quality over quantity. Better to under-deliver than to pad."
)
bullet(
    "**\"Do not ask the user clarifying questions.\"** The user has already hit submit and walked away. "
    "The agent must work with whatever it has."
)

# ============================================================
# SECTION 6 — A worked example
# ============================================================
h1("A worked example")

para(
    "Say a user with a CV full of PHP / Laravel / Vue experience searches for \"backend developer Berlin\"."
)

para("Here is what a typical run looks like in our logs:")

code(
    "[agentic-search] model=claude-sonnet-4-6 mcp=https://mcp.apify.com\n"
    "                 query=\"backend developer\" location=\"Berlin\" remote=false\n"
    "\n"
    "(Claude, thinking to itself — not visible to us — does roughly this:)\n"
    "\n"
    "  Tool call 1: apify.linkedin-jobs { query: \"Laravel developer Berlin\",\n"
    "                                     country: \"DE\", limit: 25 }\n"
    "    → 18 results. Most look relevant, a couple are senior-only.\n"
    "\n"
    "  Tool call 2: apify.indeed-jobs  { query: \"PHP backend Berlin\",\n"
    "                                    country: \"DE\", limit: 25 }\n"
    "    → 22 results. Overlaps with #1 but adds new companies.\n"
    "\n"
    "  Tool call 3: apify.glassdoor-jobs { query: \"Vue.js developer Germany\" }\n"
    "    → 14 results. A handful of frontend-leaning roles that fit the Vue side of CV.\n"
    "\n"
    "  Final: call finalize_jobs with 20 de-duplicated, relevant jobs.\n"
    "\n"
    "[agentic-search] stop_reason=tool_use mcp_calls=12 input_tokens=18450\n"
    "                 output_tokens=2140\n"
    "[agentic-search] finalized 20 jobs",
    caption="what a run looks like",
)

para(
    "Notice how Claude did not use the user's literal query (\"backend developer\"). "
    "It read the CV, saw Laravel / PHP / Vue, and searched for those — because that is what will actually match. "
    "That kind of judgment is exactly what a fixed pipeline cannot do."
)

# ============================================================
# SECTION 7 — Why this is cheap and safe
# ============================================================
h1("Cost, safety, and the things that could go wrong")

h2("Cost")
para(
    "A full agentic search costs roughly **$0.70-$1.00** in Anthropic tokens plus whatever Apify charges for the actors run "
    "(typically a few cents). That is not nothing, but it is also why we do this only once per user search, cache results in the database, "
    "and run the cheap free-tier sources alongside for free."
)

h2("Keys are the user's, not ours")
para(
    "Every user brings their own Anthropic API key and their own Apify token. "
    "We store them encrypted in the database and decrypt them per request. "
    "That means the cost sits with the user, not the app — and it also means we are not a single juicy target for key theft."
)

h2("What if Claude returns nothing useful?")
para(
    "Two safety nets: "
    "**(a)** the cheap free-tier sources run in parallel, so even if the agent returns zero we still have something to show; "
    "**(b)** the scoring pass will push junk to the bottom of the ranked list, so the user still sees the best available matches first."
)

h2("What if the agent hallucinates a job?")
para(
    "The system prompt is explicit: every finalized job must come from an actual tool result. "
    "Combined with the strict `finalize_jobs` schema (which requires a real apply URL), this is the main defense. "
    "In practice we have not seen hallucinated postings, but if we did, we would add a post-validation step that re-checks the URL."
)

# ============================================================
# SECTION 8 — Summary
# ============================================================
h1("The whole thing in 30 seconds")

bullet("**Apify** has scrapers for LinkedIn, Indeed, Glassdoor. **Claude** is an AI model. **MCP** is the plug that lets them talk.")
bullet(
    "Instead of our code calling Apify, we let **Claude** call Apify — via Apify's MCP server — picking queries, inspecting results, and refining."
)
bullet(
    "From our side, it is **one** API call to Claude with two attachments: Apify's MCP server, and a custom `finalize_jobs` tool that guarantees the final output has a clean shape."
)
bullet(
    "The agent is told to prefer quality over quantity, never invent jobs, and stop as soon as the list is strong."
)
bullet(
    "We run cheap free-tier sources in parallel as a safety net, then do a separate scoring pass to rank everything against the CV. "
    "Net effect: the user gets the kind of results Claude Desktop would give them, without the user ever having to open Claude Desktop."
)

doc.save(OUT)
print(f"wrote {OUT}")
